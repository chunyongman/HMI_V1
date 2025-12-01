# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# 제목
title = doc.add_heading('GitHub 프로젝트 클론 가이드', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# 1. 사전 준비
doc.add_heading('1. 사전 준비', level=1)
doc.add_heading('필수 프로그램 설치', level=2)
doc.add_paragraph('• Git: https://git-scm.com/downloads')
doc.add_paragraph('• Node.js: https://nodejs.org (frontend 프로젝트용)')
doc.add_paragraph('• Python: https://python.org (backend 프로젝트용)')

doc.add_heading('Git 설치 확인', level=2)
p = doc.add_paragraph()
p.add_run('git --version').bold = True

# 2. 프로젝트 클론하기
doc.add_heading('2. 프로젝트 클론하기', level=1)
doc.add_heading('기본 명령어', level=2)
p = doc.add_paragraph()
p.add_run('git clone [저장소 URL] [새 폴더명]').bold = True

doc.add_heading('ESS 프로젝트 클론 예시', level=2)
doc.add_paragraph('HMI 프로젝트:')
p = doc.add_paragraph()
p.add_run(r'git clone https://github.com/chunyongman/HMI_V1.git C:\Users\my\Desktop\새폴더\HMI_V1').italic = True

doc.add_paragraph('PLC Simulator (별도 저장소인 경우):')
p = doc.add_paragraph()
p.add_run(r'git clone https://github.com/사용자명/PLC_Simulator.git C:\Users\my\Desktop\새폴더\PLC_Simulator').italic = True

doc.add_paragraph('Edge Computer (별도 저장소인 경우):')
p = doc.add_paragraph()
p.add_run(r'git clone https://github.com/사용자명/Edge_Computer_V1.git C:\Users\my\Desktop\새폴더\Edge_Computer_V1').italic = True

# 3. 의존성 설치
doc.add_heading('3. 의존성 설치', level=1)
doc.add_heading('설치 필요 여부 확인', level=2)

table = doc.add_table(rows=4, cols=3)
table.style = 'Table Grid'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = '파일 존재 여부'
hdr_cells[1].text = '설치 명령'
hdr_cells[2].text = '설명'

row1 = table.rows[1].cells
row1[0].text = 'package.json 있음'
row1[1].text = 'npm install'
row1[2].text = 'Node.js 프로젝트'

row2 = table.rows[2].cells
row2[0].text = 'requirements.txt 있음'
row2[1].text = 'pip install -r requirements.txt'
row2[2].text = 'Python 프로젝트'

row3 = table.rows[3].cells
row3[0].text = '둘 다 없음'
row3[1].text = '설치 생략'
row3[2].text = '바로 실행 가능'

doc.add_heading('HMI_V1 프로젝트 (frontend + backend)', level=2)
doc.add_paragraph('Frontend 의존성 설치:')
p = doc.add_paragraph()
p.add_run(r'cd C:\Users\my\Desktop\새폴더\HMI_V1\frontend').italic = True
p = doc.add_paragraph()
p.add_run('npm install').italic = True

doc.add_paragraph('Backend 의존성 설치:')
p = doc.add_paragraph()
p.add_run(r'cd C:\Users\my\Desktop\새폴더\HMI_V1\backend').italic = True
p = doc.add_paragraph()
p.add_run('pip install -r requirements.txt').italic = True

# 4. 프로젝트 실행
doc.add_heading('4. 프로젝트 실행', level=1)
doc.add_heading('실행 순서 (권장)', level=2)
doc.add_paragraph('1. PLC Simulator 실행')
doc.add_paragraph('2. Edge Computer 실행')
doc.add_paragraph('3. HMI Backend 실행')
doc.add_paragraph('4. HMI Frontend 실행')

doc.add_heading('각 프로그램 실행 명령', level=2)
doc.add_paragraph('PLC Simulator:')
p = doc.add_paragraph()
p.add_run(r'cd C:\Users\my\Desktop\새폴더\PLC_Simulator').italic = True
p = doc.add_paragraph()
p.add_run('python plc_simulator.py').italic = True

doc.add_paragraph('Edge Computer:')
p = doc.add_paragraph()
p.add_run(r'cd C:\Users\my\Desktop\새폴더\Edge_Computer_V1').italic = True
p = doc.add_paragraph()
p.add_run('python main.py').italic = True

doc.add_paragraph('HMI Backend:')
p = doc.add_paragraph()
p.add_run(r'cd C:\Users\my\Desktop\새폴더\HMI_V1\backend').italic = True
p = doc.add_paragraph()
p.add_run('python main.py').italic = True

doc.add_paragraph('HMI Frontend:')
p = doc.add_paragraph()
p.add_run(r'cd C:\Users\my\Desktop\새폴더\HMI_V1\frontend').italic = True
p = doc.add_paragraph()
p.add_run('npm run dev').italic = True

# 5. 현재 저장소 URL 확인
doc.add_heading('5. 현재 저장소 URL 확인 방법', level=1)
doc.add_paragraph('기존 프로젝트의 GitHub URL을 모르는 경우:')
p = doc.add_paragraph()
p.add_run('cd [프로젝트 폴더]').italic = True
p = doc.add_paragraph()
p.add_run('git remote -v').italic = True

doc.add_paragraph('')
doc.add_paragraph('출력 예시:')
p = doc.add_paragraph()
p.add_run('origin  https://github.com/chunyongman/HMI_V1.git (fetch)').italic = True

# 6. 전체 과정 요약
doc.add_heading('6. 전체 과정 요약', level=1)
doc.add_paragraph('# 1. 새 작업 폴더 생성')
p = doc.add_paragraph()
p.add_run(r'mkdir C:\Users\my\Desktop\새프로젝트').italic = True

doc.add_paragraph('')
doc.add_paragraph('# 2. GitHub에서 클론')
p = doc.add_paragraph()
p.add_run(r'git clone https://github.com/chunyongman/HMI_V1.git C:\Users\my\Desktop\새프로젝트\HMI_V1').italic = True

doc.add_paragraph('')
doc.add_paragraph('# 3. Frontend 의존성 설치 (package.json 있으면)')
p = doc.add_paragraph()
p.add_run(r'cd C:\Users\my\Desktop\새프로젝트\HMI_V1\frontend').italic = True
p = doc.add_paragraph()
p.add_run('npm install').italic = True

doc.add_paragraph('')
doc.add_paragraph('# 4. Backend 의존성 설치 (requirements.txt 있으면)')
p = doc.add_paragraph()
p.add_run(r'cd ..\backend').italic = True
p = doc.add_paragraph()
p.add_run('pip install -r requirements.txt').italic = True

# 7. 자주 발생하는 문제
doc.add_heading('7. 자주 발생하는 문제', level=1)

doc.add_paragraph('문제: git 명령어를 찾을 수 없음', style='List Bullet')
doc.add_paragraph('    해결: Git 설치 후 터미널 재시작')

doc.add_paragraph('')
doc.add_paragraph('문제: npm install 실패', style='List Bullet')
doc.add_paragraph('    해결: Node.js 설치 확인 (node --version 으로 확인)')

doc.add_paragraph('')
doc.add_paragraph('문제: pip install 실패', style='List Bullet')
doc.add_paragraph('    해결: Python 설치 확인 (python --version 으로 확인)')

doc.add_paragraph('')
doc.add_paragraph('문제: 권한 오류', style='List Bullet')
doc.add_paragraph('    해결: 관리자 권한으로 터미널 실행')

# 문서 저장
doc.add_paragraph('')
doc.add_paragraph('문서 작성일: 2025년 12월 1일')

doc.save(r'C:\Users\my\Desktop\HMI_V1\docs\GitHub_클론_가이드.docx')
print('워드 파일 생성 완료!')
