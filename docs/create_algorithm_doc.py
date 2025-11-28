#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
인증시험 알고리즘 설명서 Word 문서 생성
"""

from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_shading(cell, color):
    """셀 배경색 설정"""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)

def add_code_block(doc, code_text):
    """코드 블록 추가"""
    p = doc.add_paragraph()
    run = p.add_run(code_text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    p.paragraph_format.left_indent = Cm(0.5)
    return p

def create_document():
    """Word 문서 생성"""
    doc = Document()

    # 제목
    title = doc.add_heading('선박용 냉각 시스템 인증 시험 알고리즘 설명서', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 문서 정보
    doc.add_paragraph()
    info = doc.add_paragraph()
    info.add_run('문서 버전: ').bold = True
    info.add_run('1.0\n')
    info.add_run('작성일: ').bold = True
    info.add_run('2025-11-28\n')
    info.add_run('대상 시스템: ').bold = True
    info.add_run('PLC Simulator + Edge Computer + HMI 통합 시스템')

    doc.add_paragraph('─' * 50)

    # 1. 시험 개요
    doc.add_heading('1. 시험 개요', level=1)
    doc.add_paragraph('본 문서는 선박용 냉각 시스템의 인증 시험에 사용되는 3가지 테스트의 알고리즘과 실행 방법을 설명합니다.')

    doc.add_heading('1.1 시험 목록', level=2)

    # 시험 목록 테이블
    table = doc.add_table(rows=4, cols=4)
    table.style = 'Table Grid'

    headers = ['시험 번호', '시험명', '측정 항목', '합격 기준']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True
        set_cell_shading(cell, 'D9E2F3')

    tests = [
        ['#1', 'PLC 제어 응답속도', 'AI 계산 완료 → VFD 변경 완료 시간', '평균 ≤0.8초, 최대 <1.0초'],
        ['#2', 'AI 예측 정확도', '온도 변화에 따른 주파수 제어 정확성', '전체 ≥90%, 항목별 ≥85%'],
        ['#3', 'HMI 실시간 반영주기', 'PLC 센서값 → HMI 화면 표시 시간', '평균 ≤1.0초'],
    ]

    for row_idx, row_data in enumerate(tests):
        for col_idx, cell_data in enumerate(row_data):
            table.rows[row_idx + 1].cells[col_idx].text = cell_data

    doc.add_paragraph()

    doc.add_heading('1.2 시스템 구성도', level=2)
    add_code_block(doc, '''┌─────────────────┐     Modbus TCP      ┌─────────────────┐
│  PLC Simulator  │◄──────────────────►│  Edge Computer  │
│   (Port 502)    │     (읽기/쓰기)     │   (AI 제어기)    │
└────────┬────────┘                     └─────────────────┘
         │
         │ Modbus TCP (읽기)
         ▼
┌─────────────────┐     WebSocket       ┌─────────────────┐
│   HMI Backend   │────────────────────►│   HMI 화면      │
│   (Port 8001)   │    (브로드캐스트)    │   (브라우저)     │
└─────────────────┘                     └─────────────────┘''')

    doc.add_paragraph('─' * 50)

    # 2. 시험 #1: PLC 제어 응답속도
    doc.add_heading('2. 시험 #1: PLC 제어 응답속도', level=1)

    doc.add_heading('2.1 시험 목적', level=2)
    doc.add_paragraph('AI 제어기가 계산을 완료한 시점부터 PLC를 통해 VFD(Variable Frequency Drive) 제어 명령이 실행 완료되기까지의 시간을 측정합니다.')

    doc.add_heading('2.2 측정 알고리즘', level=2)
    add_code_block(doc, '''┌─────────────────────────────────────────────────────────────────┐
│                    PLC 응답시간 측정 흐름                         │
├─────────────────────────────────────────────────────────────────┤
│                                                                 │
│  [1] 센서 데이터 시뮬레이션 (50개 시나리오)                       │
│      └─ 랜덤 온도/압력/부하 값 생성                              │
│                                                                 │
│  [2] AI 제어기 호출                                              │
│      └─ 온도/압력/부하 기반 최적 주파수 계산 (AI 추론 시간 측정)   │
│                                                                 │
│  [3] T1 기록 (AI 계산 완료 시점)                                  │
│                                                                 │
│  [4] PLC에 목표 주파수 쓰기 (Modbus TCP)                         │
│      └─ 10대 장비: SWP1-3, FWP1-3, FAN1-4                       │
│                                                                 │
│  [5] VFD 피드백 확인 (통신 왕복 완료 확인)                        │
│      └─ 장비 상태 읽기로 통신 성공 확인                           │
│                                                                 │
│  [6] T2 기록 (VFD 변경 완료 시점)                                │
│                                                                 │
│  [7] 응답 시간 = T2 - T1                                        │
│                                                                 │
└─────────────────────────────────────────────────────────────────┘''')

    doc.add_heading('2.3 핵심 코드 로직', level=2)
    add_code_block(doc, '''# AI 제어기 호출
sw_freq, fw_freq, fan_freq, ai_time = ai_controller.compute_optimal_frequencies(
    sensors, equipment
)

# T1: AI 계산 완료, PLC 쓰기 시작
t1 = time.perf_counter()

# PLC에 목표 주파수 쓰기
write_success = self.client.write_ai_target_frequency(target_frequencies)

# VFD 피드백 확인 (통신 왕복 완료)
while time.perf_counter() - start_wait < max_wait:
    equipment = self.client.read_equipment_status()
    if equipment:
        vfd_confirmed = True
        break
    time.sleep(0.01)  # 10ms 간격 폴링

# T2: VFD 변경 완료
t2 = time.perf_counter()
total_response_time = t2 - t1''')

    doc.add_heading('2.4 테스트 시나리오', level=2)
    table = doc.add_table(rows=4, cols=2)
    table.style = 'Table Grid'
    data = [
        ['항목', '값'],
        ['측정 횟수', '50회'],
        ['센서 범위', 'TX1-TX7: 23~48°C, PX1: 2.0~2.5bar, PU1: 30~90%'],
        ['제어 대상', 'SWP1-3, FWP1-3, FAN1-4 (총 10대)'],
    ]
    for row_idx, row_data in enumerate(data):
        for col_idx, cell_data in enumerate(row_data):
            cell = table.rows[row_idx].cells[col_idx]
            cell.text = cell_data
            if row_idx == 0:
                cell.paragraphs[0].runs[0].bold = True
                set_cell_shading(cell, 'D9E2F3')

    doc.add_paragraph()

    doc.add_heading('2.5 합격 기준', level=2)
    table = doc.add_table(rows=4, cols=3)
    table.style = 'Table Grid'
    data = [
        ['기준', '값', '설명'],
        ['평균 응답시간', '≤ 0.8초', '50회 측정 평균'],
        ['최대 응답시간', '< 1.0초', '단일 측정 최대값'],
        ['통신 성공률', '≥ 90%', '실패 횟수 10% 미만'],
    ]
    for row_idx, row_data in enumerate(data):
        for col_idx, cell_data in enumerate(row_data):
            cell = table.rows[row_idx].cells[col_idx]
            cell.text = cell_data
            if row_idx == 0:
                cell.paragraphs[0].runs[0].bold = True
                set_cell_shading(cell, 'D9E2F3')

    doc.add_paragraph()

    doc.add_heading('2.6 실행 방법', level=2)
    p = doc.add_paragraph()
    p.add_run('단계 1: PLC Simulator 시작 (별도 터미널)').bold = True
    add_code_block(doc, 'cd C:\\Users\\my\\Desktop\\PLC_Simulator\nSTART_PLC_Simulator.bat')

    p = doc.add_paragraph()
    p.add_run('단계 2: 테스트 실행').bold = True
    add_code_block(doc, 'cd C:\\Users\\my\\Desktop\\Edge_Computer_V1\npython tests/test_plc_response_time.py')

    doc.add_paragraph('필요 프로세스: PLC Simulator만 실행 (Edge Computer 모듈은 테스트에서 직접 import)')

    doc.add_paragraph('─' * 50)

    # 3. 시험 #2: AI 예측 정확도
    doc.add_heading('3. 시험 #2: AI 예측 정확도', level=1)

    doc.add_heading('3.1 시험 목적', level=2)
    doc.add_paragraph('온도 변화에 따라 AI 제어기가 올바른 방향으로 주파수를 조절하는지 검증합니다.')

    doc.add_heading('3.2 측정 알고리즘', level=2)
    add_code_block(doc, '''┌─────────────────────────────────────────────────────────────────┐
│                    AI 예측 정확도 측정 흐름                       │
├─────────────────────────────────────────────────────────────────┤
│                                                                 │
│  [Part 1] 장비별 주파수 계산 검증 (300회)                        │
│  ──────────────────────────────────────────                     │
│  각 테스트 케이스:                                               │
│    1. 온도 변화 전 상태 설정 (T_before)                          │
│    2. AI 제어기 호출 → 주파수 F1 획득                            │
│    3. 온도 변화 후 상태 설정 (T_after)                           │
│    4. AI 제어기 호출 → 주파수 F2 획득                            │
│    5. 주파수 변화 방향 판단:                                     │
│       - F2 > F1 + 0.5Hz → "UP" (증가)                           │
│       - F2 < F1 - 0.5Hz → "DOWN" (감소)                         │
│       - 그 외 → "STABLE" (유지)                                 │
│    6. 기대 방향과 실제 방향 비교                                 │
│                                                                 │
│  [Part 2] 온도 예측 방향성 검증 (90회)                           │
│  ──────────────────────────────────────────                     │
│  각 테스트 케이스:                                               │
│    1. 5개 온도 데이터 포인트 생성 (T_start → T_end 선형 보간)    │
│    2. 온도 추세 예측 알고리즘 실행:                              │
│       - 최근 3개 데이터의 평균 변화량 계산                        │
│       - 변화량 > 0.3°C → "UP"                                   │
│       - 변화량 < -0.3°C → "DOWN"                                │
│       - 그 외 → "STABLE"                                        │
│    3. 기대 방향과 예측 방향 비교                                 │
│                                                                 │
└─────────────────────────────────────────────────────────────────┘''')

    doc.add_heading('3.3 제어 로직 원리', level=2)

    p = doc.add_paragraph()
    p.add_run('3.3.1 SWP (Sea Water Pump) 제어').bold = True
    add_code_block(doc, '''제어 센서: T5 (CLR FW Out Temp)
목표: FW 출구 온도 유지

  T5 > 40°C  →  주파수 증가 (60Hz) - 냉각수 순환 강화
  T5 < 30°C  →  주파수 감소 (40Hz) - 과냉각 방지
  30-40°C    →  피드백 제어 (점진적 조절)''')

    p = doc.add_paragraph()
    p.add_run('3.3.2 FWP (Fresh Water Pump) 제어').bold = True
    add_code_block(doc, '''제어 센서: T4 (CLR FW In Temp)
목표: FW 입구 온도 관리

  T4 ≥ 48°C  →  주파수 증가 (60Hz) - 긴급 냉각
  T4 < 38°C  →  주파수 감소 (40Hz) - 과냉각 방지
  38-48°C    →  비례 제어''')

    p = doc.add_paragraph()
    p.add_run('3.3.3 FAN (Engine Room Fan) 제어').bold = True
    add_code_block(doc, '''제어 센서: T6 (E/R Inside Temp)
목표: 기관실 온도 43°C 유지

  T6 ≥ 47°C  →  주파수 증가 (60Hz) - 긴급 배기
  T6 < 43°C  →  피드백 제어 (목표 온도 추종)
  43-47°C    →  PID 피드백 제어''')

    doc.add_heading('3.4 테스트 케이스 구성', level=2)
    table = doc.add_table(rows=10, cols=4)
    table.style = 'Table Grid'
    data = [
        ['카테고리', '테스트 수', '온도 범위', '설명'],
        ['SWP 상승', '50회', 'T5: 35→41°C+', '온도↑ → 주파수↑'],
        ['SWP 하강', '50회', 'T5: 32→29°C-', '온도↓ → 주파수↓'],
        ['FWP 상승', '50회', 'T4: 44→48°C+', '온도↑ → 주파수↑'],
        ['FWP 하강', '50회', 'T4: 42→37°C-', '온도↓ → 주파수↓'],
        ['FAN 상승', '50회', 'T6: 43→47°C+', '온도↑ → 주파수↑'],
        ['FAN 하강', '50회', 'T6: 48→43°C-', '온도↓ → 주파수↓'],
        ['예측(SWP)', '30회', '상승10+하강10+안정10', '방향성 예측'],
        ['예측(FWP)', '30회', '상승10+하강10+안정10', '방향성 예측'],
        ['예측(FAN)', '30회', '상승10+하강10+안정10', '방향성 예측'],
    ]
    for row_idx, row_data in enumerate(data):
        for col_idx, cell_data in enumerate(row_data):
            cell = table.rows[row_idx].cells[col_idx]
            cell.text = cell_data
            if row_idx == 0:
                cell.paragraphs[0].runs[0].bold = True
                set_cell_shading(cell, 'D9E2F3')

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('총 390회 테스트').bold = True

    doc.add_heading('3.5 합격 기준', level=2)
    table = doc.add_table(rows=6, cols=3)
    table.style = 'Table Grid'
    data = [
        ['기준', '값', '설명'],
        ['전체 정확도', '≥ 90%', '390회 중 351회 이상 성공'],
        ['SWP 정확도', '≥ 85%', '100회 중 85회 이상 성공'],
        ['FWP 정확도', '≥ 85%', '100회 중 85회 이상 성공'],
        ['FAN 정확도', '≥ 85%', '100회 중 85회 이상 성공'],
        ['예측 정확도', '≥ 85%', '90회 중 77회 이상 성공'],
    ]
    for row_idx, row_data in enumerate(data):
        for col_idx, cell_data in enumerate(row_data):
            cell = table.rows[row_idx].cells[col_idx]
            cell.text = cell_data
            if row_idx == 0:
                cell.paragraphs[0].runs[0].bold = True
                set_cell_shading(cell, 'D9E2F3')

    doc.add_paragraph()

    doc.add_heading('3.6 실행 방법', level=2)
    doc.add_paragraph('※ 이 테스트는 외부 프로세스 연결 없이 단독 실행 가능합니다.')
    add_code_block(doc, 'cd C:\\Users\\my\\Desktop\\Edge_Computer_V1\npython tests/test_ai_prediction_accuracy.py')
    doc.add_paragraph('필요 프로세스: 없음 (AI 제어 모듈만 import하여 순수 계산 로직 검증)')

    doc.add_paragraph('─' * 50)

    # 4. 시험 #3: HMI 실시간 반영주기
    doc.add_heading('4. 시험 #3: HMI 실시간 반영주기', level=1)

    doc.add_heading('4.1 시험 목적', level=2)
    doc.add_paragraph('PLC의 센서값이 변경된 후 HMI 화면에 반영되기까지의 시간을 측정합니다.')

    doc.add_heading('4.2 측정 알고리즘', level=2)
    add_code_block(doc, '''┌─────────────────────────────────────────────────────────────────┐
│                    HMI 반영주기 측정 흐름                         │
├─────────────────────────────────────────────────────────────────┤
│                                                                 │
│  [측정 원리: Value Match 방식]                                   │
│  ────────────────────────────                                   │
│                                                                 │
│  PLC Simulator는 1초마다 자체적으로 센서값을 변동시킴             │
│  이 자연스러운 변동을 활용하여 데이터 전파 시간을 측정             │
│                                                                 │
│  [1] PLC에서 현재 센서값 읽기 (T1 기록)                          │
│      └─ Modbus TCP로 레지스터 10-16 읽기                         │
│      └─ TX1~TX7 값 획득                                         │
│                                                                 │
│  [2] HMI WebSocket에서 같은 값 대기                              │
│      └─ WebSocket 메시지 모니터링                                │
│      └─ 허용 오차: ±0.15 (반올림 차이)                           │
│                                                                 │
│  [3] 일치하는 값 수신 시점 (T2 기록)                             │
│                                                                 │
│  [4] 반영 시간 = T2 - T1                                        │
│                                                                 │
│  [데이터 흐름]                                                   │
│  PLC Simulator ──Modbus TCP──► HMI Backend ──WebSocket──► 테스트 │
│  (센서값 생성)    (1초 주기 읽기)   (1초 주기 브로드캐스트)        │
│                                                                 │
└─────────────────────────────────────────────────────────────────┘''')

    doc.add_heading('4.3 핵심 코드 로직', level=2)
    add_code_block(doc, '''def measure_refresh_time(self, test_id: int) -> RefreshMeasurement:
    # 1. 테스트할 센서 선택 (TX1~TX7 순환)
    sensor = self.SENSORS[(test_id - 1) % len(self.SENSORS)]

    # 2. PLC에서 현재값 읽기 (T1)
    plc_read_time = time.time()
    plc_sensors = self.plc_client.read_all_sensors()
    plc_value = plc_sensors[sensor]

    # 3. HMI WebSocket에서 같은 값 수신 대기 (T2)
    hmi_time = self.ws_receiver.wait_for_value(
        sensor_key=sensor,
        expected_value=plc_value,
        timeout=2.0,
        tolerance=0.15  # ±0.15 허용 (반올림 오차)
    )

    # 4. 반영 시간 계산
    refresh_time = hmi_time - plc_read_time

    return RefreshMeasurement(
        refresh_time=refresh_time,
        passed=(refresh_time <= 1.0)
    )''')

    doc.add_heading('4.4 테스트 구성', level=2)
    table = doc.add_table(rows=6, cols=2)
    table.style = 'Table Grid'
    data = [
        ['항목', '값'],
        ['측정 횟수', '100회'],
        ['측정 센서', 'TX1~TX7 (순환)'],
        ['타임아웃', '2.0초'],
        ['허용 오차', '±0.15 (반올림 차이)'],
        ['측정 간격', '0.15초'],
    ]
    for row_idx, row_data in enumerate(data):
        for col_idx, cell_data in enumerate(row_data):
            cell = table.rows[row_idx].cells[col_idx]
            cell.text = cell_data
            if row_idx == 0:
                cell.paragraphs[0].runs[0].bold = True
                set_cell_shading(cell, 'D9E2F3')

    doc.add_paragraph()

    doc.add_heading('4.5 합격 기준', level=2)
    p = doc.add_paragraph()
    p.add_run('평균 반영 시간 ≤ 1.0초').bold = True
    p.add_run(' (100회 측정 평균)')

    doc.add_heading('4.6 실행 방법', level=2)
    p = doc.add_paragraph()
    p.add_run('단계 1: PLC Simulator 시작 (별도 터미널)').bold = True
    add_code_block(doc, 'cd C:\\Users\\my\\Desktop\\PLC_Simulator\nSTART_PLC_Simulator.bat')

    p = doc.add_paragraph()
    p.add_run('단계 2: HMI Backend 시작 (별도 터미널)').bold = True
    add_code_block(doc, 'cd C:\\Users\\my\\Desktop\\HMI_V1\nSTART_HMI_V1.bat')

    p = doc.add_paragraph()
    p.add_run('단계 3: 테스트 실행 (새 터미널)').bold = True
    add_code_block(doc, 'cd C:\\Users\\my\\Desktop\\HMI_V1\npython tests/test_hmi_refresh_rate.py')

    doc.add_paragraph('필요 프로세스: PLC Simulator + HMI Backend')

    doc.add_paragraph('─' * 50)

    # 5. 시험별 필요 프로세스 요약
    doc.add_heading('5. 시험별 필요 프로세스 요약', level=1)

    table = doc.add_table(rows=4, cols=4)
    table.style = 'Table Grid'
    headers = ['시험', 'PLC Simulator', 'Edge Computer', 'HMI Backend']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True
        set_cell_shading(cell, 'D9E2F3')

    processes = [
        ['#1 PLC 응답속도', 'O', '모듈만 import', 'X'],
        ['#2 AI 예측 정확도', 'X', '모듈만 import', 'X'],
        ['#3 HMI 반영주기', 'O', 'X', 'O'],
    ]

    for row_idx, row_data in enumerate(processes):
        for col_idx, cell_data in enumerate(row_data):
            cell = table.rows[row_idx + 1].cells[col_idx]
            cell.text = cell_data
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    doc.add_paragraph('─' * 50)

    # 6. 전체 시험 실행 순서
    doc.add_heading('6. 전체 시험 실행 순서', level=1)

    add_code_block(doc, '''┌────────────────────────────────────────────────────────────────┐
│  1. PLC Simulator 시작                                         │
│     cd C:\\Users\\my\\Desktop\\PLC_Simulator                     │
│     START_PLC_Simulator.bat                                    │
│     (5초 대기)                                                  │
├────────────────────────────────────────────────────────────────┤
│  2. 시험 #1: PLC 응답시간                                       │
│     cd C:\\Users\\my\\Desktop\\Edge_Computer_V1                  │
│     python tests/test_plc_response_time.py                     │
├────────────────────────────────────────────────────────────────┤
│  3. 시험 #2: AI 예측 정확도 (단독 실행)                          │
│     cd C:\\Users\\my\\Desktop\\Edge_Computer_V1                  │
│     python tests/test_ai_prediction_accuracy.py                │
├────────────────────────────────────────────────────────────────┤
│  4. 시험 #3: HMI 실시간 반영주기 (HMI 추가 실행 필요)            │
│     cd C:\\Users\\my\\Desktop\\HMI_V1                            │
│     START_HMI_V1.bat  (새 터미널에서)                           │
│     python tests/test_hmi_refresh_rate.py                      │
└────────────────────────────────────────────────────────────────┘''')

    doc.add_paragraph('─' * 50)

    # 7. 기술 사양
    doc.add_heading('7. 기술 사양', level=1)

    doc.add_heading('7.1 통신 프로토콜', level=2)
    table = doc.add_table(rows=4, cols=3)
    table.style = 'Table Grid'
    data = [
        ['구간', '프로토콜', '포트'],
        ['Edge Computer ↔ PLC', 'Modbus TCP', '502'],
        ['HMI Backend ↔ PLC', 'Modbus TCP', '502'],
        ['HMI Backend → 브라우저', 'WebSocket', '8001'],
    ]
    for row_idx, row_data in enumerate(data):
        for col_idx, cell_data in enumerate(row_data):
            cell = table.rows[row_idx].cells[col_idx]
            cell.text = cell_data
            if row_idx == 0:
                cell.paragraphs[0].runs[0].bold = True
                set_cell_shading(cell, 'D9E2F3')

    doc.add_paragraph()

    doc.add_heading('7.2 데이터 갱신 주기', level=2)
    table = doc.add_table(rows=4, cols=2)
    table.style = 'Table Grid'
    data = [
        ['구성요소', '갱신 주기'],
        ['PLC Simulator 센서값', '1초'],
        ['HMI Backend PLC 읽기', '1초'],
        ['HMI WebSocket 브로드캐스트', '1초'],
    ]
    for row_idx, row_data in enumerate(data):
        for col_idx, cell_data in enumerate(row_data):
            cell = table.rows[row_idx].cells[col_idx]
            cell.text = cell_data
            if row_idx == 0:
                cell.paragraphs[0].runs[0].bold = True
                set_cell_shading(cell, 'D9E2F3')

    doc.add_paragraph()

    doc.add_heading('7.3 센서 레지스터 매핑', level=2)
    table = doc.add_table(rows=10, cols=3)
    table.style = 'Table Grid'
    data = [
        ['센서', '레지스터', '설명'],
        ['TX1', '10', 'CSW PP Disc Temp'],
        ['TX2', '11', 'No.1 CLR SW Out Temp'],
        ['TX3', '12', 'No.2 CLR SW Out Temp'],
        ['TX4', '13', 'CLR FW In Temp'],
        ['TX5', '14', 'CLR FW Out Temp'],
        ['TX6', '15', 'E/R Inside Temp'],
        ['TX7', '16', 'E/R Outside Temp'],
        ['PX1', '17', 'CSW PP Disc Press'],
        ['PU1', '19', 'M/E Load'],
    ]
    for row_idx, row_data in enumerate(data):
        for col_idx, cell_data in enumerate(row_data):
            cell = table.rows[row_idx].cells[col_idx]
            cell.text = cell_data
            if row_idx == 0:
                cell.paragraphs[0].runs[0].bold = True
                set_cell_shading(cell, 'D9E2F3')

    doc.add_paragraph()
    doc.add_paragraph('─' * 50)
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('문서 끝').bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 저장
    output_path = r'C:\Users\my\Desktop\HMI_V1\docs\인증시험_알고리즘_설명서.docx'
    doc.save(output_path)
    print(f"Word 문서 생성 완료: {output_path}")
    return output_path

if __name__ == "__main__":
    create_document()
