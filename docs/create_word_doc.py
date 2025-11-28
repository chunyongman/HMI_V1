#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""인증시험 테스트 가이드 Word 문서 생성"""

from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_shading(cell, color):
    """셀 배경색 설정"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), color)
    tcPr.append(shd)

def create_document():
    doc = Document()

    # 기본 스타일 설정
    style = doc.styles['Normal']
    style.font.name = '맑은 고딕'
    style.font.size = Pt(10)

    # 제목
    title = doc.add_heading('PLC 시뮬레이터 & Edge Computer 인증 시험 가이드', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 문서 정보
    doc.add_paragraph()
    info = doc.add_paragraph()
    info.add_run('문서 정보').bold = True
    doc.add_paragraph('• 작성일: 2025년 11월 28일')
    doc.add_paragraph('• 버전: 1.0')
    doc.add_paragraph('• 목적: 인증기관 제출용 시험 절차 및 방법 안내')

    # 구분선
    doc.add_paragraph('─' * 50)

    # 1. 시험 개요
    doc.add_heading('1. 시험 개요', level=1)

    doc.add_heading('1.1 시험 목적', level=2)
    doc.add_paragraph('PLC 시뮬레이터와 Edge Computer 시스템의 성능 및 신뢰성을 검증하기 위한 인증 시험입니다.')

    doc.add_heading('1.2 시험 항목 요약', level=2)

    # 시험 항목 테이블
    table = doc.add_table(rows=4, cols=4)
    table.style = 'Table Grid'

    # 헤더
    headers = ['시험 번호', '시험 항목', '합격 기준', '테스트 파일 위치']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True
        set_cell_shading(cell, 'D9E2F3')

    # 데이터
    data = [
        ['#1', 'PLC 응답시간', '평균 <= 1.0초', 'Edge_Computer_V1/tests/test_plc_response_time.py'],
        ['#2', 'AI 예측 정확도', '정확도 >= 90%', 'Edge_Computer_V1/tests/test_ai_prediction_accuracy.py'],
        ['#3', 'HMI 실시간 반영주기', '평균 <= 1.0초', 'HMI_V1/tests/test_hmi_refresh_rate.py'],
    ]

    for row_idx, row_data in enumerate(data):
        for col_idx, cell_data in enumerate(row_data):
            table.rows[row_idx + 1].cells[col_idx].text = cell_data

    doc.add_paragraph()

    doc.add_heading('1.3 시스템 구성도', level=2)
    p = doc.add_paragraph()
    p.add_run('''
┌─────────────────┐     Modbus TCP      ┌─────────────────┐
│  PLC Simulator  │◄──────────────────►│  Edge Computer  │
│   (Port 502)    │                     │                 │
└─────────────────┘                     └─────────────────┘
                                                │
                                                │ WebSocket
                                                ▼
                                        ┌─────────────────┐
                                        │    HMI (Web)    │
                                        └─────────────────┘
''').font.name = 'Consolas'

    # 구분선
    doc.add_paragraph('─' * 50)

    # 2. 사전 준비사항
    doc.add_heading('2. 사전 준비사항', level=1)

    doc.add_heading('2.1 필수 소프트웨어', level=2)
    doc.add_paragraph('• Python 3.10 이상')
    doc.add_paragraph('• PLC 시뮬레이터 (Modbus TCP 지원)')
    doc.add_paragraph('• 필요 Python 패키지: pymodbus, numpy, pandas')

    doc.add_heading('2.2 폴더 구조', level=2)
    p = doc.add_paragraph()
    p.add_run('''
C:\\Users\\my\\Desktop\\
├── PLC_Simulator\\
│   ├── plc_simulator.py
│   └── START_PLC_Simulator.bat            (PLC 시뮬레이터 시작)
│
├── Edge_Computer_V1\\
│   ├── tests\\
│   │   ├── test_plc_response_time.py      (시험 #1)
│   │   └── test_ai_prediction_accuracy.py (시험 #2)
│   └── test_results\\                      (결과 저장)
│
└── HMI_V1\\
    ├── tests\\
    │   └── test_hmi_refresh_rate.py       (시험 #3)
    ├── test_results\\                      (결과 저장)
    └── START_HMI.bat                       (HMI 시작)
''').font.name = 'Consolas'

    # 구분선
    doc.add_paragraph('─' * 50)

    # 3. 시험 #1
    doc.add_heading('3. 시험 #1: PLC 응답시간 검증', level=1)

    doc.add_heading('3.1 시험 목적', level=2)
    doc.add_paragraph('PLC 시뮬레이터와 Edge Computer 간 Modbus TCP 통신의 응답 시간을 측정합니다.')

    doc.add_heading('3.2 합격 기준', level=2)
    p = doc.add_paragraph()
    p.add_run('평균 응답시간 <= 1.0초 (1초 이내)').bold = True

    doc.add_heading('3.3 테스트 원리', level=2)
    doc.add_paragraph('1. Edge Computer → PLC: 레지스터 읽기 요청 (T1 기록)')
    doc.add_paragraph('2. PLC → Edge Computer: 응답 수신 (T2 기록)')
    doc.add_paragraph('3. 응답시간 = T2 - T1')
    doc.add_paragraph('4. 100회 반복 측정하여 평균 계산')

    doc.add_heading('3.4 실행 방법', level=2)

    p = doc.add_paragraph()
    p.add_run('단계 1: PLC 시뮬레이터 시작').bold = True
    p = doc.add_paragraph()
    p.add_run('cd C:\\Users\\my\\Desktop\\PLC_Simulator\nSTART_PLC_Simulator.bat').font.name = 'Consolas'

    p = doc.add_paragraph()
    p.add_run('단계 2: 테스트 실행').bold = True
    p = doc.add_paragraph()
    p.add_run('cd C:\\Users\\my\\Desktop\\Edge_Computer_V1\npython tests/test_plc_response_time.py').font.name = 'Consolas'

    p = doc.add_paragraph()
    p.add_run('단계 3: Enter 키를 눌러 시험 시작').bold = True

    doc.add_heading('3.5 결과 파일', level=2)
    doc.add_paragraph('• 상세 결과: Edge_Computer_V1/test_results/test_results_plc_response_YYYYMMDD_HHMMSS.csv')
    doc.add_paragraph('• 요약 결과: Edge_Computer_V1/test_results/test_summary_plc_response_YYYYMMDD_HHMMSS.csv')

    # 구분선
    doc.add_paragraph('─' * 50)

    # 4. 시험 #2
    doc.add_heading('4. 시험 #2: AI 예측 정확도 검증', level=1)

    doc.add_heading('4.1 시험 목적', level=2)
    doc.add_paragraph('Edge Computer의 AI 기반 제어 시스템이 다양한 센서 입력 조건에서 올바른 VFD 주파수를 예측하는지 검증합니다.')

    doc.add_heading('4.2 합격 기준', level=2)
    p = doc.add_paragraph()
    p.add_run('정확도 >= 90%').bold = True
    doc.add_paragraph('(500개 테스트 케이스 중 450개 이상 통과)')

    doc.add_heading('4.3 테스트 원리', level=2)
    doc.add_paragraph('1. 미리 정의된 500개 테스트 시나리오 실행')
    doc.add_paragraph('2. 각 시나리오: 센서값 → AI 예측 → 예상 결과와 비교')
    doc.add_paragraph('3. 허용 오차: ±1Hz')
    doc.add_paragraph('4. 정확도 = (통과 케이스 / 전체 케이스) × 100%')

    doc.add_heading('4.4 테스트 시나리오 구성', level=2)

    # 시나리오 테이블
    table = doc.add_table(rows=11, cols=3)
    table.style = 'Table Grid'

    headers = ['시나리오', '설명', '케이스 수']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True
        set_cell_shading(cell, 'D9E2F3')

    scenarios = [
        ['T1', '해수온도 상승 (냉각 부하 증가)', '50개'],
        ['T2', '해수온도 하강 (냉각 부하 감소)', '50개'],
        ['T3', 'SW Outlet 온도 상승', '50개'],
        ['T4', 'SW Outlet 온도 하강', '50개'],
        ['T5', 'FW Outlet 온도 상승', '50개'],
        ['T6', 'FW Outlet 온도 하강', '50개'],
        ['T7', 'Engine Room 온도 상승', '50개'],
        ['T8', 'Engine Room 온도 하강', '50개'],
        ['T9', '엔진 부하 증가', '50개'],
        ['T10', '엔진 부하 감소', '50개'],
    ]

    for row_idx, row_data in enumerate(scenarios):
        for col_idx, cell_data in enumerate(row_data):
            table.rows[row_idx + 1].cells[col_idx].text = cell_data

    doc.add_paragraph()

    doc.add_heading('4.5 실행 방법', level=2)

    doc.add_paragraph('※ 이 테스트는 외부 프로세스 연결 없이 단독 실행 가능합니다.')
    doc.add_paragraph('(Edge Computer의 AI 제어 모듈만 import하여 순수 계산 로직 검증)')
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.add_run('단계 1: 테스트 실행').bold = True
    p = doc.add_paragraph()
    p.add_run('cd C:\\Users\\my\\Desktop\\Edge_Computer_V1\npython tests/test_ai_prediction_accuracy.py').font.name = 'Consolas'

    p = doc.add_paragraph()
    p.add_run('단계 2: Enter 키를 눌러 시험 시작').bold = True

    doc.add_heading('4.6 결과 파일', level=2)
    doc.add_paragraph('• 상세 결과: Edge_Computer_V1/test_results/test_results_ai_prediction_YYYYMMDD_HHMMSS.csv')
    doc.add_paragraph('• 요약 결과: Edge_Computer_V1/test_results/test_summary_ai_prediction_YYYYMMDD_HHMMSS.csv')

    # 구분선
    doc.add_paragraph('─' * 50)

    # 5. 시험 #3
    doc.add_heading('5. 시험 #3: HMI 실시간 반영주기 검증', level=1)

    doc.add_heading('5.1 시험 목적', level=2)
    doc.add_paragraph('PLC 센서값 변경 후 HMI 화면에 반영되기까지의 시간을 측정합니다.')

    doc.add_heading('5.2 합격 기준', level=2)
    p = doc.add_paragraph()
    p.add_run('평균 반영 주기 <= 1.0초 (1초 이내)').bold = True

    doc.add_heading('5.3 테스트 원리', level=2)
    doc.add_paragraph('1. PLC Simulator에 새로운 센서값 쓰기 (T1 기록)')
    doc.add_paragraph('2. HMI WebSocket에서 변경된 값 감지 대기')
    doc.add_paragraph('3. HMI에서 값 변경 감지 시점 (T2 기록)')
    doc.add_paragraph('4. 반영 시간 = T2 - T1')
    doc.add_paragraph('5. 100회 반복 측정하여 평균 계산')
    doc.add_paragraph()
    doc.add_paragraph('측정 경로: PLC Simulator -> HMI WebSocket -> 화면 표시')

    doc.add_heading('5.4 측정 대상 센서', level=2)

    # 센서 테이블
    table = doc.add_table(rows=10, cols=3)
    table.style = 'Table Grid'

    headers = ['센서', '레지스터', '설명']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True
        set_cell_shading(cell, 'D9E2F3')

    sensors = [
        ['TX1', '10', 'CSW PP Disc Temp'],
        ['TX2', '11', 'No.1 CLR SW Out Temp'],
        ['TX3', '12', 'No.2 CLR SW Out Temp'],
        ['TX4', '13', 'CLR FW In Temp'],
        ['TX5', '14', 'CLR FW Out Temp'],
        ['TX6', '15', 'E/R Inside Temp'],
        ['TX7', '16', 'E/R Outside Temp'],
        ['PX1', '17', 'CSW PP Disc Press'],
        ['PU1', '19', 'M/E Load'],
    ]

    for row_idx, row_data in enumerate(sensors):
        for col_idx, cell_data in enumerate(row_data):
            table.rows[row_idx + 1].cells[col_idx].text = cell_data

    doc.add_paragraph()

    doc.add_heading('5.5 실행 방법', level=2)

    p = doc.add_paragraph()
    p.add_run('단계 1: PLC 시뮬레이터 시작').bold = True
    p = doc.add_paragraph()
    p.add_run('cd C:\\Users\\my\\Desktop\\PLC_Simulator\nSTART_PLC_Simulator.bat').font.name = 'Consolas'

    p = doc.add_paragraph()
    p.add_run('단계 2: HMI Backend 시작 (새 터미널에서)').bold = True
    p = doc.add_paragraph()
    p.add_run('cd C:\\Users\\my\\Desktop\\HMI_V1\nSTART_HMI.bat').font.name = 'Consolas'
    doc.add_paragraph('※ WebSocket 서버가 localhost:8001에서 실행됨')

    p = doc.add_paragraph()
    p.add_run('단계 3: 테스트 실행 (새 터미널에서)').bold = True
    p = doc.add_paragraph()
    p.add_run('cd C:\\Users\\my\\Desktop\\HMI_V1\npython tests/test_hmi_refresh_rate.py').font.name = 'Consolas'

    p = doc.add_paragraph()
    p.add_run('단계 4: Enter 키를 눌러 시험 시작').bold = True

    doc.add_heading('5.6 결과 파일', level=2)
    doc.add_paragraph('• 상세 결과: HMI_V1/test_results/test_results_hmi_refresh_YYYYMMDD_HHMMSS.csv')
    doc.add_paragraph('• 요약 결과: HMI_V1/test_results/test_summary_hmi_refresh_YYYYMMDD_HHMMSS.csv')

    # 구분선
    doc.add_paragraph('─' * 50)

    # 6. 전체 시험 실행 순서 요약
    doc.add_heading('6. 전체 시험 실행 순서 요약', level=1)

    doc.add_heading('6.1 권장 실행 순서', level=2)

    p = doc.add_paragraph()
    p.add_run('''
┌────────────────────────────────────────────────────────────────┐
│  1. PLC 시뮬레이터 시작                                         │
│     cd C:\\Users\\my\\Desktop\\PLC_Simulator                     │
│     START_PLC_Simulator.bat                                    │
│     (5초 대기)                                                  │
├────────────────────────────────────────────────────────────────┤
│  2. 시험 #1: PLC 응답시간                                       │
│     cd C:\\Users\\my\\Desktop\\Edge_Computer_V1                  │
│     python tests/test_plc_response_time.py                     │
├────────────────────────────────────────────────────────────────┤
│  3. 시험 #2: AI 예측 정확도 (단독 실행)                          │
│     cd C:\\Users\\my\\Desktop\\Edge_Computer_V1                  │
│     python tests/test_ai_prediction_accuracy.py                │
├────────────────────────────────────────────────────────────────┤
│  4. 시험 #3: HMI 실시간 반영주기 (HMI 추가 실행 필요)           │
│     cd C:\\Users\\my\\Desktop\\HMI_V1                            │
│     START_HMI.bat  (새 터미널에서)                              │
│     python tests/test_hmi_refresh_rate.py                      │
└────────────────────────────────────────────────────────────────┘
''').font.name = 'Consolas'

    doc.add_heading('6.2 시험별 필요 프로세스', level=2)

    # 프로세스 테이블
    table = doc.add_table(rows=4, cols=4)
    table.style = 'Table Grid'

    headers = ['시험', 'PLC Simulator', 'Edge Computer', 'HMI']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True
        set_cell_shading(cell, 'D9E2F3')

    processes = [
        ['#1 PLC 응답시간', 'O', 'O (모듈만)', 'X'],
        ['#2 AI 예측 정확도', 'X', 'O (모듈만)', 'X'],
        ['#3 HMI 반영주기', 'O', 'X', 'O'],
    ]

    for row_idx, row_data in enumerate(processes):
        for col_idx, cell_data in enumerate(row_data):
            cell = table.rows[row_idx + 1].cells[col_idx]
            cell.text = cell_data
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    doc.add_paragraph('참고:')
    doc.add_paragraph('• 시험 #1: PLC Simulator 필요, Edge Computer 모듈만 import (별도 실행 불필요)')
    doc.add_paragraph('• 시험 #2: 단독 실행 가능 (PLC 연결 없이 AI 계산 로직만 검증)')
    doc.add_paragraph('• 시험 #3: PLC Simulator + HMI Backend 필요 (실제 화면 반영 측정)')

    # 구분선
    doc.add_paragraph()
    doc.add_paragraph('─' * 50)

    # 7. 문제 해결
    doc.add_heading('7. 문제 해결', level=1)

    doc.add_heading('7.1 PLC 연결 실패', level=2)
    p = doc.add_paragraph()
    p.add_run('[ERROR] PLC 시뮬레이터에 연결할 수 없습니다.').font.name = 'Consolas'

    p = doc.add_paragraph()
    p.add_run('해결방법:').bold = True
    doc.add_paragraph('1. PLC 시뮬레이터가 실행 중인지 확인')
    doc.add_paragraph('2. START_HMI_V1.bat 실행 후 5초 이상 대기')
    doc.add_paragraph('3. 포트 502가 사용 중인지 확인')

    doc.add_heading('7.2 모듈 임포트 오류', level=2)
    p = doc.add_paragraph()
    p.add_run("ModuleNotFoundError: No module named 'pymodbus'").font.name = 'Consolas'

    p = doc.add_paragraph()
    p.add_run('해결방법:').bold = True
    p = doc.add_paragraph()
    p.add_run('pip install pymodbus==2.5.3\npip install numpy pandas').font.name = 'Consolas'

    doc.add_heading('7.3 결과 파일 저장 실패', level=2)
    p = doc.add_paragraph()
    p.add_run('해결방법:').bold = True
    doc.add_paragraph('• test_results 폴더가 존재하는지 확인')
    doc.add_paragraph('• 폴더 쓰기 권한 확인')

    # 구분선
    doc.add_paragraph()
    doc.add_paragraph('─' * 50)

    # 문서 끝
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('─ 문서 끝 ─').bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 저장
    output_path = r'C:\Users\my\Desktop\HMI_V1\docs\인증시험_테스트_가이드_final.docx'
    doc.save(output_path)
    print(f"Word 문서 생성 완료: {output_path}")
    return output_path

if __name__ == "__main__":
    create_document()
